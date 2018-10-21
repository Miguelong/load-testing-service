#!/usr/bin/python
# -*- coding: UTF-8 -*-

import Queue
import json
import os
import threading
import time

import MySQLdb
import requests
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import pylab as pl
from docx import Document
from docx.shared import Pt


class LoadTest:
    # temp dir where generated report and images are placed
    TEMP_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'temp')

    # dictionary store the stop flags of tests, the key is test_id
    STOP_SETTING = {}

    def __init__(self, url, concurrent_num, method, header, payload, timeout, proxy, parameters_list, test_id):
        self.url = url
        self.concurrent_num = concurrent_num
        self.method = method
        self.header = header
        self.payload = payload
        self.timeout = timeout
        self.proxy = proxy
        self.parameters_list = parameters_list
        self.test_id = test_id
        # load test begins
        self.begin_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        # load test starts
        self.start = time.time()
        # tmsp is used in the file names
        self.tmsp = time.strftime("%Y%m%d%H%M%S", time.localtime())
        # self.stop = False
        self.STOP_SETTING[test_id] = False

    def statistic(self, title, queue, document):
        dict = {}
        sum = 0.0
        count = 0
        max = 0.0
        min = 100.0
        data = []
        response_time_trend = []
        if queue.empty():
            return

        document.add_heading(title + ': ' + str(queue.qsize()), 2)

        # count the sum of requests of each response time
        # get the trend of response time
        while not queue.empty():
            element = queue.get()
            total_seconds = round(element[1], 2)
            response_time_trend.append(element)
            data.append(total_seconds)
            if dict.has_key(total_seconds):
                dict[total_seconds] = dict[total_seconds] + 1
            else:
                dict[total_seconds] = 1

            if total_seconds > max:
                max = total_seconds

            if total_seconds < min:
                min = total_seconds

            sum += total_seconds
            count = count + 1

        average = round(sum / count, 3)

        items = dict.items()
        items.sort()

        # draw figure of X-'Response time' Y-'Count'
        bins = int((max - min) / 0.01)
        if bins == 0:
            bins = 1

        plt.hist(data, bins=bins, color='steelblue', edgecolor='k', label='title')
        plt.xlabel('Response time (s)')
        plt.ylabel('Count')

        img = str(self.test_id) + "_" + self.tmsp + '_' + title + ".png"
        plt.savefig(os.path.join(self.TEMP_DIR, img))
        plt.close()
        # add figure to document
        document.add_picture(os.path.join(self.TEMP_DIR, img))

        # draw figure of X-'Request start' Y-'Response time'
        response_time_trend.sort()
        x = []
        y = []
        for ele in response_time_trend:
            x.append(ele[0])
            y.append(ele[1])
        pl.plot(x, y, '.')
        pl.xlabel('Request start (s)')
        pl.ylabel('Response time (s)')

        img = str(self.test_id) + "_" + self.tmsp + '_' + title + "_trend.png"
        pl.savefig(os.path.join(self.TEMP_DIR, img))
        pl.close()
        # add figure to document
        document.add_picture(os.path.join(self.TEMP_DIR, img))

        paragraph = document.add_paragraph('')
        run = paragraph.add_run('max:' + str(max) + ' min:' + str(min) + ' average:' + str(average) + '\n\n')
        run.font.size = Pt(12)

        left_bound = min
        right_bound = min + 0.5
        interval_count = 0
        detail = ""
        run = paragraph.add_run('[' + str(left_bound) + ',' + str(right_bound) + '): ')
        run.font.size = Pt(12)
        for item in items:
            total_seconds = item[0]
            count = item[1]
            while total_seconds >= right_bound:
                left_bound = right_bound
                right_bound = right_bound + 0.5
                if total_seconds < right_bound:
                    if len(detail) > 0:
                        detail = detail[0:len(detail) - 1]
                    run = paragraph.add_run(str(interval_count)+'\n')
                    run.font.size = Pt(12)
                    run = paragraph.add_run(detail + '\n\n')
                    run.font.size = Pt(12)
                    run = paragraph.add_run('[' + str(left_bound) + ',' + str(right_bound) + '): ')
                    run.font.size = Pt(12)
                    interval_count = 0
                    detail = ""

            detail = detail + "(" + str(total_seconds) + "," + str(count) + "),"
            interval_count = interval_count + count

        if len(detail) > 0:
            detail = detail[0:len(detail) - 1]
            run = paragraph.add_run(str(interval_count) + '\n' + detail + '\n')
            run.font.size = Pt(12)

    def output_fails(self, queue, document):
        if queue.empty():
            return

        document.add_heading('Fail ', 1)
        paragraph = document.add_paragraph('')
        run = paragraph.add_run('total:' + str(queue.qsize()))
        run.font.size = Pt(12)

        dict = {}
        while not queue.empty():
            item = queue.get()
            key = item[0]
            value = item[1]
            if dict.has_key(key):
                dict[key].append(value)
            else:
                dict[key] = [value]

        items = dict.items()
        for item in items:
            # print >> file, item[0]
            document.add_heading(item[0], 2)
            # print >> file, item[1]
            paragraph = document.add_paragraph('')
            run = paragraph.add_run('total:' + str(len(item[1])) + '\n')
            run.font.size = Pt(12)
            run = paragraph.add_run(item[1])
            run.font.size = Pt(12)
            run = paragraph.add_run('\n\n')
            run.font.size = Pt(12)

    def request_get(self, url):
        request_start = round(time.time()-self.start, 2)
        # request_start = time.time()-self.start
        try:
            # r = requests.get(url, timeout=timeout)
            print url
            r = requests.get(url)
            json = r.json()
            print json
            code = json['meta']['code']
            total_seconds = r.elapsed.total_seconds()
            return None, code, total_seconds, request_start
        except Exception, e:
            return e, None, None

    def request_post(self, url, payload):
        request_start = round(time.time() - self.start, 2)
        # request_start = time.time() - self.start
        try:
            # r = requests.post(url, json=payload, timeout=timeout)
            r = requests.post(url, json=payload)
            json = r.json()
            code = json['meta']['code']
            total_seconds = r.elapsed.total_seconds()
            return None, code, total_seconds,request_start
        except Exception, e:
            return e, None, None

    def request_get_proxy(self, url):
        request_start = round(time.time() - self.start, 2)
        # request_start = time.time() - self.start
        try:
            if self.header is None:
                cmd = "curl -w 'time_total: %{time_total}\n' '" + url + "' -H 'Content-Type:application/json' --proxy " + self.proxy
            else:
                cmd = "curl -w 'time_total: %{time_total}\n' '" + url + "' -H 'Content-Type:application/json' -H '" + self.header + "' --proxy " + self.proxy

            res = os.popen(cmd).read()
            index = res.find('time_total')
            code = json.loads(res[0:index])['meta']['code']
            total_seconds = float(res[index + 12:].replace("\n", ""))
            return None, code, total_seconds, request_start
        except Exception, e:
            return e, None, None

    def request_post_proxy(self, url, payload):
        request_start = round(time.time() - self.start, 2)
        # request_start = time.time() - self.start
        try:
            if self.header is None:
                cmd = "curl -w 'time_total: %{time_total}\n' '" + url + "' -X POST -H 'Content-Type:application/json' -d '" + payload + "' --proxy " + self.proxy
            else:
                cmd = "curl -w 'time_total: %{time_total}\n' '" + url + "' -X POST -H 'Content-Type:application/json' -H '" + self.header + "' -d '" + payload + "' --proxy " + self.proxy

            res = os.popen(cmd).read()
            index = res.find('time_total')
            code = json.loads(res[0:index])['meta']['code']
            total_seconds = float(res[index + 12:].replace("\n", ""))
            return None, code, total_seconds, request_start
        except Exception, e:
            return e, None, None

    def process_response(self, res, timeout, successQueue, timeoutQueue, failQueue, parameters):
        e = res[0]
        code = res[1]
        total_seconds = res[2]
        if e is None:
            request_start = res[3]
            if code == 200:
                if total_seconds > timeout:
                    timeoutQueue.put((request_start, total_seconds))
                else:
                    successQueue.put((request_start, total_seconds))
            else:
                failQueue.put(("code " + str(code), parameters))
        else:
            failQueue.put((repr(e), parameters))

    def testGet(self, url_template, timeout, successQueue, timeoutQueue, failQueue, parametersList):
        # get count of parameters of the url
        parameters_count = url_template.count('${')
        for parameters in parametersList:
            # in case of user stop the test
            # if self.stop:
            if self.STOP_SETTING[self.test_id]:
                break

            parameter_list = json.loads(parameters)
            url = url_template
            for i in range(parameters_count):
                old = '${' + str(i) + '}'
                new = parameter_list[i]
                url = url.replace(old, new)
            if self.proxy is None:
                res = self.request_get(url)
            else:
                res = self.request_get_proxy(url)
            self.process_response(res, timeout, successQueue, timeoutQueue, failQueue, parameters)

    def testPost(self, url, payload_template, timeout, successQueue, timeoutQueue, failQueue, parametersList):
        # get count of parameters of the payload
        parameters_count = payload_template.count('${')
        for parameters in parametersList:
            # in case of user stop the test
            # if self.stop:
            if self.STOP_SETTING[self.test_id]:
                break

            list = json.loads(parameters)
            payload = payload_template
            for i in range(parameters_count):
                old = '${' + str(i) + '}'
                new = list[i]
                payload = payload.replace(old, new)

            if self.proxy is None:
                res = self.request_post(url, json.loads(payload))
            else:
                res = self.request_post_proxy(url, payload)

            self.process_response(res, timeout, successQueue, timeoutQueue, failQueue, parameters)

    def get_progress(self, total_num, successQueue, timeoutQueue, failQueue):
        db = MySQLdb.connect("model-mysql.internal.gridx.com", "demo", "RE3u6pc8ZYx1c", "test")
        cursor = db.cursor()

        progress = 0.0
        while progress < 100.0:
            # in case of user stop the test
            # if self.stop:
            if self.STOP_SETTING[self.test_id]:
                break

            time.sleep(3)
            processed_num = successQueue.qsize() + timeoutQueue.qsize() + failQueue.qsize()
            progress = float(processed_num) / total_num * 100
            print total_num, processed_num
            print progress
            try:
                cursor.execute("update load_test set progress=" + str(progress) + " where id=" + str(self.test_id))
                db.commit()
            except Exception, e:
                print repr(e)
                db.rollback()

        if progress == 100.0:
            status = 'completed'
        else:
            status = 'stopped'

        try:
            cursor.execute("update load_test set progress="+str(progress)+", status='"+status+"' where id=" + str(self.test_id))
            db.commit()
        except Exception, e:
            print repr(e)
            db.rollback()

        db.close()

        self.get_report(successQueue, timeoutQueue, failQueue)

    # get report of load test
    def get_report(self, successQueue, timeoutQueue, failQueue, ):
        end_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        document = Document()

        # its begin time and end time
        document.add_heading('Report', 0)
        paragraph = document.add_paragraph('')
        run = paragraph.add_run('begin: ' + self.begin_time + '\n')
        run.font.size = Pt(12)
        run = paragraph.add_run('end: ' + end_time)
        run.font.size = Pt(12)

        # testing parameters
        document.add_heading('Testing Parameters', 1)
        paragraph = document.add_paragraph('')
        run = paragraph.add_run('Concurrent threads: ' + str(self.concurrent_num) + '\n')
        run.font.size = Pt(12)
        run = paragraph.add_run('Url: ' + self.url + '\n')
        run.font.size = Pt(12)
        if self.payload is not None:
            run = paragraph.add_run('Payload: ' + self.payload + '\n')
            run.font.size = Pt(12)
        if self.header is not None:
            run = paragraph.add_run('Header: ' + self.header + '\n')
            run.font.size = Pt(12)
        if self.proxy is not None:
            run = paragraph.add_run('Proxy: ' + self.proxy + '\n')
            run.font.size = Pt(12)

        # statistics of test
        document.add_heading('Statistics', 1)
        self.statistic('Success', successQueue, document)
        self.statistic('Timeout', timeoutQueue, document)

        # output fails
        self.output_fails(failQueue, document)

        # set the file name and save it to temp dir
        doc = str(self.test_id) + '_' + self.tmsp + '.docx'
        document.save(os.path.join(self.TEMP_DIR, doc))

        # update field 'report' of the record
        try:
            db = MySQLdb.connect("model-mysql.internal.gridx.com", "demo", "RE3u6pc8ZYx1c", "test")
            cursor = db.cursor()
            sql = "update load_test set report='%s' where id='%s'" % \
                  (doc, str(self.test_id))
            cursor.execute(sql)
            db.commit()
        except Exception,e:
            print repr(e)
            db.rollback()

        db.close()

    def start_test(self):
        # allocate parameters for each thread
        parameters = []
        begin = 0
        total_num = len(self.parameters_list)

        if total_num < self.concurrent_num:
            size = 1
            end = size
            for i in range(self.concurrent_num):
                if i < total_num:
                    arr = self.parameters_list[begin:end]
                    begin = begin + size
                    end = end + size
                else:
                    arr = []
                parameters.append(arr)
        else:
            size = total_num / self.concurrent_num
            end = size
            for i in range(self.concurrent_num):
                if i == self.concurrent_num - 1:
                    end = total_num

                arr = self.parameters_list[begin:end]
                parameters.append(arr)
                begin = begin + size
                end = end + size

        successQueue = Queue.Queue(total_num)
        timeoutQueue = Queue.Queue(total_num)
        failQueue = Queue.Queue(total_num)
        threads = []
        for i in range(self.concurrent_num):
            parametersPerThread = parameters[i]
            if self.method == 'get':
                t = threading.Thread(target=self.testGet, args=(
                self.url, self.timeout, successQueue, timeoutQueue, failQueue, parametersPerThread,))
            elif self.method == 'post':
                t = threading.Thread(target=self.testPost, args=(
                self.url, self.payload, self.timeout, successQueue, timeoutQueue, failQueue, parametersPerThread,))

            t.setDaemon(True)
            t.start()
            threads.append(t)

        t = threading.Thread(target=self.get_progress, args=(total_num, successQueue, timeoutQueue, failQueue,))
        t.setDaemon(True)
        t.start()
        threads.append(t)

    @classmethod
    def stop_test(cls, test_id):
        cls.STOP_SETTING[test_id] = True
